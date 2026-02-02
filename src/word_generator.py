from docx import Document
from docx.shared import Cm, Inches
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os
import logging

logger = logging.getLogger(__name__)

class WordGenerator:
    """處理 Word 文件排版與產出的類別"""

    def __init__(self):
        # A4 標準尺寸 (21cm x 29.7cm)
        self.PAGE_WIDTH_CM = 21.0
        self.MARGIN_CM = 1.27  # 窄邊界預設值

    def calculate_max_columns(self, qr_size_cm, page_width_cm=21.0, margin_cm=1.27):
        """計算一列最多可容納幾個 QR Code"""
        available_width = page_width_cm - (margin_cm * 2)
        # 額外扣除一點微小的間隙 (0.1cm) 以確保不會擠出去
        cols = int(available_width / (qr_size_cm + 0.1))
        return max(1, cols)

    def create_report(self, data_list, qr_size_cm, output_path):
        """
        建立包含 QR Code 的 Word 文件
        :param data_list: 包含 {"text": str, "img_path": str} 的列表
        :param qr_size_cm: QR Code 的尺寸 (cm)
        :param output_path: 輸出的路徑
        """
        logger.info(f"開始建立 Word 文件: {output_path}")
        doc = Document()

        # 設定 A4 紙張與邊界
        section = doc.sections[0]
        section.page_height = Cm(29.7)
        section.page_width = Cm(21.0)
        section.left_margin = Cm(self.MARGIN_CM)
        section.right_margin = Cm(self.MARGIN_CM)
        section.top_margin = Cm(self.MARGIN_CM)
        section.bottom_margin = Cm(self.MARGIN_CM)

        num_cols = self.calculate_max_columns(qr_size_cm, self.PAGE_WIDTH_CM, self.MARGIN_CM)
        num_items = len(data_list)
        # 每一項佔用兩橫列：一列放圖片，一列放文字
        num_rows = ((num_items + num_cols - 1) // num_cols) * 2

        table = doc.add_table(rows=num_rows, cols=num_cols)
        table.alignment = WD_ALIGN_PARAGRAPH.CENTER

        for i, item in enumerate(data_list):
            row_idx = (i // num_cols) * 2
            col_idx = i % num_cols
            
            # 第一列放圖片
            cell_img = table.cell(row_idx, col_idx)
            paragraph_img = cell_img.paragraphs[0]
            paragraph_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = paragraph_img.add_run()
            run.add_picture(item["img_path"], width=Cm(qr_size_cm))
            
            # 第二列放文字標籤
            cell_text = table.cell(row_idx + 1, col_idx)
            cell_text.text = item["text"]
            cell_text.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 設定儲存格垂直置中
            cell_img.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            cell_text.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 確保目錄存在
        os.makedirs(os.path.dirname(os.path.abspath(output_path)), exist_ok=True)
        doc.save(output_path)
        logger.info(f"Word 文件儲存成功: {output_path}")
